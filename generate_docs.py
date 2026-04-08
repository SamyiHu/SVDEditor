#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""生成SVD Editor两份独立文档：使用说明书 + 维护文档"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '2E75B6')
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(value)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
            if r % 2 == 1:
                set_cell_shading(cell, 'D9E2F3')
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    return table

def setup_font(doc):
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_code_block(doc, code, size=Pt(9)):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = size
    p.paragraph_format.left_indent = Cm(0.5)
    return p

def add_tip(doc, title, text, color=RGBColor(0x1F, 0x49, 0x7D)):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = color
    p.add_run(text)

# ============================================================
#  文档一：使用说明书
# ============================================================
def generate_user_manual():
    doc = Document()
    setup_font(doc)

    # 封面
    for _ in range(5):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run('SVD Editor 使用说明书')
    run.font.size = Pt(32)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run('SVD文件可视化编辑工具\n版本 V1.0\n2026年4月')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    doc.add_page_break()

    # --- 第1章 软件简介 ---
    doc.add_heading('第1章 软件简介', level=1)
    doc.add_paragraph(
        'SVD Editor是一款SVD（System View Description）文件可视化编辑工具。'
        'SVD文件是ARM CMSIS标准定义的XML格式文件，用于描述微控制器的外设、寄存器、位域和中断信息。'
        '本软件将这些XML信息以图形界面方式呈现，让用户可以直观地创建和编辑SVD文件。'
    )
    doc.add_heading('1.1 功能概览', level=2)
    add_styled_table(doc,
        ['功能', '说明'],
        [
            ['基本信息编辑', '芯片名称、版本、厂商、版权、CPU配置等'],
            ['外设管理', '外设的增删改、继承（derivedFrom）、排序'],
            ['寄存器管理', '寄存器偏移、大小、访问权限、复位值'],
            ['位域管理', '位域偏移、位宽、访问权限'],
            ['位域图', '寄存器位域可视化分布图，悬停查看详情，点击选中编辑'],
            ['中断管理', '多外设共用中断，搜索+全选快捷操作'],
            ['排序操作', '外设按名称/地址排序，寄存器按偏移排序'],
            ['连锁操作', '删除元素时自动联动删除关联元素，支持通配符和变量替换'],
            ['实时预览', '编辑时实时生成SVD XML预览'],
            ['文件操作', '打开/保存/另存为SVD文件'],
            ['撤销/重做', '多步撤销与重做'],
        ],
        col_widths=[4, 12]
    )

    # --- 第2章 环境搭建 ---
    doc.add_heading('第2章 环境搭建', level=1)
    doc.add_heading('2.1 系统要求', level=2)
    add_styled_table(doc,
        ['项目', '要求'],
        [
            ['操作系统', 'Windows 10/11（推荐）、macOS 12+、Ubuntu 20.04+'],
            ['Python', '3.9 或更高版本'],
            ['内存', '≥ 4GB RAM'],
            ['磁盘', '≥ 200MB 可用空间'],
            ['显示器', '推荐分辨率 ≥ 1280×720'],
        ],
        col_widths=[4, 12]
    )

    doc.add_heading('2.2 从源码运行', level=2)
    doc.add_paragraph('步骤一：获取源代码')
    add_code_block(doc, 'git clone https://github.com/SamyiHu/SVDEditor.git\ncd SVDEditor')
    doc.add_paragraph('步骤二：安装Python依赖')
    add_code_block(doc, 'pip install -r requirements.txt')
    doc.add_paragraph('依赖说明：')
    add_styled_table(doc,
        ['依赖包', '版本', '用途'],
        [
            ['PyQt6', '≥ 6.5.0', 'GUI框架（必须）'],
            ['pytest', '≥ 7.0.0', '单元测试（可选）'],
            ['pytest-qt', '≥ 4.0.0', 'GUI测试（可选）'],
            ['python-docx', '≥ 1.0.0', '文档生成（可选）'],
        ],
        col_widths=[4, 3, 8]
    )
    doc.add_paragraph('步骤三：启动程序')
    add_code_block(doc, 'python run.py')

    doc.add_heading('2.3 打包为可执行文件', level=2)
    doc.add_paragraph('如需生成独立的 .exe 可执行文件：')
    add_code_block(doc, 'python build.py')
    doc.add_paragraph('打包后的文件在 dist/ 目录中，可直接双击运行，无需安装Python环境。')

    # --- 第3章 界面介绍 ---
    doc.add_heading('第3章 界面介绍', level=1)
    doc.add_paragraph(
        '软件主界面分为四个区域：\n\n'
        '① 左侧面板 — 外设树：以树形结构展示所有外设、寄存器和位域。\n'
        '   点击外设节点可查看外设详情，点击寄存器/位域节点可切换到对应编辑页面。\n\n'
        '② 右上区域 — 标签页编辑区：包含以下标签页：\n'
        '   · 基本信息：芯片名称、版本、厂商、CPU配置\n'
        '   · 外设：当前选中外设的属性编辑\n'
        '   · 寄存器：当前选中寄存器的属性编辑\n'
        '   · 位域：当前选中位域的属性编辑\n'
        '   · 中断：所有中断的管理列表\n\n'
        '③ 右下区域 — XML预览：实时显示生成的SVD XML内容。\n\n'
        '④ 底部状态栏：显示当前操作状态和数据统计（外设数/寄存器数/位域数/中断数）。'
    )

    # --- 第4章 快速上手 ---
    doc.add_heading('第4章 快速上手', level=1)

    doc.add_heading('4.1 典型工作流程', level=2)
    doc.add_paragraph('以下是最常见的使用流程：')
    doc.add_paragraph()
    doc.add_paragraph('【流程一：编辑现有SVD文件】')
    doc.add_paragraph('1. 文件 → 打开（Ctrl+O），选择 .svd 文件')
    doc.add_paragraph('2. 软件自动解析并展示所有信息')
    doc.add_paragraph('3. 在左侧树中点击要编辑的外设/寄存器/位域')
    doc.add_paragraph('4. 在右侧编辑区修改属性')
    doc.add_paragraph('5. 文件 → 保存（Ctrl+S）或 另存为（Ctrl+Shift+S）')
    doc.add_paragraph()
    doc.add_paragraph('【流程二：从零创建SVD文件】')
    doc.add_paragraph('1. 文件 → 新建（Ctrl+N）')
    doc.add_paragraph('2. 在"基本信息"标签页填写芯片名称、版本等')
    doc.add_paragraph('3. 在"外设"标签页添加外设（点击"添加外设"按钮）')
    doc.add_paragraph('4. 选中外设后，在"寄存器"标签页添加寄存器')
    doc.add_paragraph('5. 选中寄存器后，在"位域"标签页添加位域')
    doc.add_paragraph('6. 在"中断"标签页管理中断')
    doc.add_paragraph('7. 保存文件')

    doc.add_heading('4.2 快捷键', level=2)
    add_styled_table(doc,
        ['快捷键', '功能'],
        [
            ['Ctrl+O', '打开SVD文件'],
            ['Ctrl+S', '保存'],
            ['Ctrl+Shift+S', '另存为'],
            ['Ctrl+Z', '撤销'],
            ['Ctrl+Y', '重做'],
            ['Delete', '删除选中项'],
            ['双击列表项', '编辑'],
        ],
        col_widths=[4, 10]
    )

    # --- 第5章 功能详解 ---
    doc.add_heading('第5章 功能详解', level=1)

    # 5.1 基本信息
    doc.add_heading('5.1 基本信息', level=2)
    doc.add_paragraph('在"基本信息"标签页中设置芯片的顶层描述：')
    add_styled_table(doc,
        ['字段', '说明', '是否必填'],
        [
            ['芯片名称', 'MCU型号，如 SC32R807TS8', '是'],
            ['描述', '芯片功能描述', '否'],
            ['版本号', '文件版本，如 1.0', '是'],
            ['SVD版本', 'SVD格式版本，通常为 1.3', '是'],
            ['厂商名称', '厂商名称。勾选"不显示"则不会在SVD中输出<vendor>元素', '否'],
            ['版权信息', '版权声明。勾选"不显示"则不输出', '否'],
            ['作者', '文件维护者姓名', '否'],
            ['许可证', '选择开源许可证类型', '否'],
        ],
        col_widths=[3, 8, 2]
    )
    doc.add_paragraph()
    add_tip(doc, '💡 提示：', '厂商名称、版权信息、作者和许可证在打开SVD文件时会自动解析并回填到对应字段中。')

    doc.add_heading('CPU配置区域', level=3)
    doc.add_paragraph('在基本信息页面下方有CPU配置区域：')
    doc.add_paragraph('• CPU名称：选择内核类型（CM0/CM3/CM4/CM7等）', style='List Bullet')
    doc.add_paragraph('• 字节序：小端（little）或大端（big）', style='List Bullet')
    doc.add_paragraph('• MPU/FPU：是否支持内存保护单元和浮点运算单元', style='List Bullet')
    doc.add_paragraph('• NVIC优先级位数：中断优先级位数（通常为2或4）', style='List Bullet')

    # 5.2 外设管理
    doc.add_heading('5.2 外设管理', level=2)
    doc.add_paragraph('外设是SVD文件的核心组织单元。每个外设包含一组寄存器和可选的中断定义。')

    doc.add_heading('5.2.1 添加外设', level=3)
    doc.add_paragraph('1. 切换到"外设"标签页')
    doc.add_paragraph('2. 点击"添加外设"按钮')
    doc.add_paragraph('3. 在弹出的对话框中填写：')
    add_styled_table(doc,
        ['属性', '说明'],
        [
            ['名称', '外设标识符，如 GPIO、UART0、TIM1（必须唯一）'],
            ['基地址', '外设在地址空间中的起始地址，如 0x40000000'],
            ['显示名称', '用于显示的友好名称（可选）'],
            ['描述', '外设功能描述（可选）'],
            ['组名', '外设分组名称，如 GPIO组、Timer组（可选）'],
            ['地址块偏移', '默认 0x0'],
            ['地址块大小', '地址范围大小，如 0x100'],
            ['继承自', '选择要继承的基础外设（可选）'],
        ],
        col_widths=[3, 12]
    )

    doc.add_heading('5.2.2 编辑外设', level=3)
    doc.add_paragraph('在左侧树中点击外设节点，右侧会显示该外设的属性。点击"编辑"按钮修改。')

    doc.add_heading('5.2.3 删除外设', level=3)
    doc.add_paragraph('选中要删除的外设，点击"删除"按钮。注意：删除外设会同时删除其下的所有寄存器和位域。')

    doc.add_heading('5.2.4 外设继承', level=3)
    doc.add_paragraph(
        'SVD支持外设继承（derivedFrom）功能。当多个外设具有相同的寄存器结构时，'
        '可以只在一个"基础外设"中定义寄存器，其他外设通过继承自动获得相同的寄存器定义。\n\n'
        '使用方法：在添加/编辑外设时，在"继承自"下拉框中选择基础外设。'
    )

    doc.add_heading('5.2.5 外设排序', level=3)
    doc.add_paragraph('• 按名称排序：将所有外设按字母顺序排列', style='List Bullet')
    doc.add_paragraph('• 按地址排序：将所有外设按基地址从小到大排列', style='List Bullet')

    # 5.3 寄存器管理
    doc.add_heading('5.3 寄存器管理', level=2)
    doc.add_paragraph('选中某个外设后，在"寄存器"标签页中管理该外设的寄存器。')
    add_styled_table(doc,
        ['属性', '说明'],
        [
            ['名称', '寄存器名称，如 CTRL、DATA、STATUS'],
            ['偏移地址', '相对于外设基地址的偏移，如 0x00、0x04'],
            ['大小', '寄存器位宽，如 0x20（32位）、0x10（16位）、0x08（8位）'],
            ['访问权限', 'read-write / read-only / write-only'],
            ['复位值', '上电后的默认值，如 0x00000000'],
            ['描述', '寄存器功能说明'],
        ],
        col_widths=[3, 12]
    )

    # 5.4 位域管理
    doc.add_heading('5.4 位域管理', level=2)
    doc.add_paragraph('选中某个寄存器后，在"位域"标签页中管理该寄存器的位域。')
    add_styled_table(doc,
        ['属性', '说明'],
        [
            ['名称', '位域名称，如 EN、MODE、IRQ_FLAG'],
            ['起始位', '位域在寄存器中的起始位置（从0开始）'],
            ['位宽', '位域占据的位数（1~32）'],
            ['访问权限', 'read-write / read-only / write-only'],
            ['复位值', '该位域的默认值'],
        ],
        col_widths=[3, 12]
    )
    add_tip(doc, '💡 示例：', '一个32位寄存器的 [7:0] 位域，起始位=0，位宽=8。[31:16] 位域，起始位=16，位宽=16。')

    # 5.5 中断管理
    doc.add_heading('5.5 中断管理', level=2)
    doc.add_paragraph(
        '中断管理是本软件的重要特性之一。在实际MCU中，经常出现多个外设共用同一个中断向量的情况。'
        '例如 UART0、UART2、UART4 共用中断号16。SVD Editor完整支持这种多对多关系。'
    )

    doc.add_heading('5.5.1 查看中断列表', level=3)
    doc.add_paragraph(
        '切换到"中断"标签页，以表格形式展示所有中断。每行包含：\n'
        '· 名称 — 中断名称\n'
        '· 中断号 — 中断向量编号\n'
        '· 关联外设 — 所有共享此中断的外设（逗号分隔）\n'
        '· 描述 — 中断描述'
    )

    doc.add_heading('5.5.2 添加中断', level=3)
    doc.add_paragraph('1. 点击"添加中断"按钮')
    doc.add_paragraph('2. 在弹出的对话框中填写中断名称和中断号')
    doc.add_paragraph('3. 在"关联外设"区域勾选需要关联的外设')
    doc.add_paragraph('4. 点击确定')
    doc.add_paragraph()
    add_tip(doc, '快捷操作：', '')
    doc.add_paragraph('• 搜索框：输入关键字（如"UART"）过滤外设列表', style='List Bullet')
    doc.add_paragraph('• 全选：一键选中所有当前可见的外设', style='List Bullet')
    doc.add_paragraph('• 清空：一键取消选中所有当前可见的外设', style='List Bullet')
    doc.add_paragraph('• 已选计数：实时显示已选中多少个外设', style='List Bullet')
    doc.add_paragraph()
    add_tip(doc, '💡 使用技巧：', '要快速选中所有UART外设 → 在搜索框输入"UART" → 点击"全选" → 完成！')

    doc.add_heading('5.5.3 编辑中断', level=3)
    doc.add_paragraph('在中断列表中选中要编辑的中断，点击"编辑"按钮。可以修改名称、中断号、增减关联外设。')

    doc.add_heading('5.5.4 删除中断', level=3)
    doc.add_paragraph('选中要删除的中断，点击"删除"。注意：如果该中断关联了多个外设，会从所有外设中同时移除。')

    # 5.6 位域图
    doc.add_heading('5.6 寄存器位域图', level=2)
    doc.add_paragraph(
        '在选中一个寄存器后，软件会自动在位域标签页上方显示该寄存器的位域分布图。'
        '这是一个可视化的32位寄存器位域示意图，直观展示各个位域的分布和属性。'
    )

    doc.add_heading('5.6.1 位域图功能', level=3)
    doc.add_paragraph('• 彩色色块：每个位域用不同颜色标识，方便区分', style='List Bullet')
    doc.add_paragraph('• 位域名称：直接在色块中显示位域名称', style='List Bullet')
    doc.add_paragraph('• 保留位：未定义的位用灰色标注为 "RESERVED"', style='List Bullet')
    doc.add_paragraph('• 位号标尺：顶部显示完整的32位编号标尺', style='List Bullet')
    doc.add_paragraph('• 悬停提示：鼠标悬停在位域上时，显示详细信息（名称、偏移、位宽、访问权限、复位值）', style='List Bullet')
    doc.add_paragraph('• 点击选中：点击位域可选中对应的位域进行编辑', style='List Bullet')
    doc.add_paragraph('• 跳转外设：对于继承的位域，点击后可跳转到源外设查看', style='List Bullet')

    doc.add_heading('5.6.2 位域图颜色说明', level=3)
    add_styled_table(doc,
        ['颜色区域', '含义'],
        [
            ['蓝色系/绿色系/暖色系色块', '已定义的位域，不同颜色用于区分不同位域'],
            ['灰色区域', '保留位（RESERVED），未定义的位'],
        ],
        col_widths=[5, 11]
    )
    add_tip(doc, '💡 提示：', '位域图是实时更新的。编辑位域后，位域图会自动刷新。对于继承来的寄存器，位域图中也能看到基础外设的位域定义。')

    # 5.7 排序操作
    doc.add_heading('5.7 排序操作', level=2)
    doc.add_paragraph(
        'SVD Editor提供多种排序功能，帮助用户整理和组织SVD文件中的数据。'
    )

    doc.add_heading('5.7.1 外设排序', level=3)
    doc.add_paragraph('在外设标签页中，提供两种排序方式：')
    doc.add_paragraph('• 按名称排序：将所有外设按字母顺序（A-Z）排列', style='List Bullet')
    doc.add_paragraph('• 按地址排序：将所有外设按基地址从小到大排列', style='List Bullet')
    doc.add_paragraph()
    add_tip(doc, '💡 使用场景：', '当SVD文件中的外设顺序比较混乱时，使用排序功能可以快速整理。排序不会影响数据内容，只改变显示和输出顺序。')

    doc.add_heading('5.7.2 寄存器排序', level=3)
    doc.add_paragraph('寄存器按偏移地址自动排序显示，确保偏移地址小的寄存器排在前面。')

    doc.add_heading('5.7.3 中断排序', level=3)
    doc.add_paragraph('中断列表默认按中断号排序，方便查看和查找特定中断。')

    # 5.8 连锁操作
    doc.add_heading('5.8 连锁操作', level=2)
    doc.add_paragraph(
        '连锁操作是SVD Editor的一项高级功能，可以在删除某个元素时自动联动删除/修改其他关联元素。'
        '这在MCU外设之间存在依赖关系时非常有用。'
    )
    add_tip(doc, '典型场景：', 'GPIO外设的端口引脚（如PA0）与配置寄存器（PACON.MODE0）、上下拉寄存器（PAPH.PUPD0）、电平寄存器（PALEV.LEV0）存在对应关系。删除PA0时，应该同步删除所有相关位域。')

    doc.add_heading('5.8.1 启用/关闭连锁功能', level=3)
    doc.add_paragraph('菜单栏 → 工具 → 启用连锁操作，可以全局开关连锁功能。')
    doc.add_paragraph('菜单栏 → 工具 → 编辑连锁规则...，可以打开规则编辑界面。')

    doc.add_heading('5.8.2 连锁规则编辑', level=3)
    doc.add_paragraph('在连锁规则编辑界面中，可以添加、编辑和删除规则。每条规则包含：')
    add_styled_table(doc,
        ['字段', '说明', '示例'],
        [
            ['规则名称', '标识这条规则', '"GPIO-端口配置连锁"'],
            ['源类型', '触发源的层级', 'peripheral / register / field'],
            ['源外设', '触发源的外设名（支持通配符）', '"GPIOA" / "GPIO*" / "*"'],
            ['源寄存器', '触发源的寄存器名', '"MODER" / "*"'],
            ['源位域', '触发源的位域名', '"PA0" / "PA*"'],
            ['触发条件', '何时触发', 'delete / modify'],
            ['动作列表', '触发后执行的操作', '见下文'],
        ],
        col_widths=[3, 4, 8]
    )

    doc.add_heading('5.8.3 通配符说明', level=3)
    add_styled_table(doc,
        ['通配符', '含义', '示例'],
        [
            ['*', '匹配所有', '"*" → 匹配任何名称'],
            ['前缀*', '匹配前缀', '"GPIO*" → 匹配 GPIOA、GPIOB、GPIOC...'],
            ['*后缀', '匹配后缀', '"*CON" → 匹配 PACON、PBCON...'],
            ['精确值', '完全匹配', '"GPIOA" → 只匹配GPIOA'],
        ],
        col_widths=[3, 3, 9]
    )

    doc.add_heading('5.8.4 变量替换', level=3)
    doc.add_paragraph('在动作目标中可以使用变量，运行时自动替换为实际值：')
    add_styled_table(doc,
        ['变量', '替换为', '示例'],
        [
            ['$PERIPHERAL', '源外设名称', '触发源是GPIOA → 替换为"GPIOA"'],
            ['$REGISTER', '源寄存器名称', '触发源是MODER → 替换为"MODER"'],
            ['$FIELD', '源位域名称', '触发源是PA0 → 替换为"PA0"'],
        ],
        col_widths=[3, 4, 8]
    )

    doc.add_heading('5.8.5 连锁操作示例', level=3)
    doc.add_paragraph('【示例1】删除 GPIOA.PA.PA0 → 同步删除三个关联位域：')
    doc.add_paragraph('· 自动删除 PACON.MODE0（端口配置）', style='List Bullet')
    doc.add_paragraph('· 自动删除 PAPH.PUPD0（上下拉配置）', style='List Bullet')
    doc.add_paragraph('· 自动删除 PALEV.LEV0（电平控制）', style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('【示例2】使用通配符实现批量连锁：')
    doc.add_paragraph('设置源为 GPIO*（所有GPIO外设），删除任何GPIO位域时，自动删除对应外设的CON/PH/LEV寄存器中的相关位域。')

    # 5.9 实时预览
    doc.add_heading('5.9 实时预览', level=2)
    doc.add_paragraph(
        '软件右下角的预览区域会实时显示当前SVD文件的XML内容。每次编辑操作后，预览会自动更新。'
        '也可以点击"预览XML"按钮手动刷新。\n\n'
        '预览区域显示的是标准的SVD XML格式，可以直接复制使用。'
    )

    # 5.10 文件操作
    doc.add_heading('5.10 文件操作', level=2)
    add_styled_table(doc,
        ['操作', '方式', '说明'],
        [
            ['打开文件', '文件 → 打开（Ctrl+O）', '支持 .svd 和 .xml 文件'],
            ['保存文件', '文件 → 保存（Ctrl+S）', '保存到当前文件路径'],
            ['另存为', '文件 → 另存为（Ctrl+Shift+S）', '选择新的保存路径'],
            ['新建文件', '文件 → 新建（Ctrl+N）', '创建空白SVD'],
        ],
        col_widths=[3, 5, 7]
    )

    # --- 第6章 常见问题 ---
    doc.add_heading('第6章 常见问题', level=1)
    add_styled_table(doc,
        ['问题', '解决方案'],
        [
            ['打开文件后厂商/版权信息为空', '确认SVD文件包含<vendor>元素和注释中的Copyright信息。新版本已支持自动解析。'],
            ['中断只关联了一个外设', '在编辑中断对话框中勾选多个外设即可。可使用搜索+全选快速操作。'],
            ['继承的外设无法编辑寄存器', '继承来的寄存器是只读的。如需修改，请在基础外设中修改。'],
            ['中文乱码', '确保SVD文件使用UTF-8编码。'],
            ['程序启动报错', '运行 pip install -r requirements.txt 安装依赖。'],
        ],
        col_widths=[5, 11]
    )

    # 保存
    path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'SVD_Editor_使用说明书.docx')
    doc.save(path)
    print(f'使用说明书已生成: {path}')
    return path


# ============================================================
#  文档二：维护文档
# ============================================================
def generate_maintenance_doc():
    doc = Document()
    setup_font(doc)

    # 封面
    for _ in range(5):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run('SVD Editor 维护文档')
    run.font.size = Pt(32)
    run.bold = True
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run('代码结构 · 数据模型 · 扩展指南\n版本 V1.0\n2026年4月')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    doc.add_page_break()

    # --- 第1章 项目架构 ---
    doc.add_heading('第1章 项目架构', level=1)

    doc.add_heading('1.1 目录结构', level=2)
    add_code_block(doc, '''SVDEditor/
├── run.py                          # 程序入口
├── build.py                        # PyInstaller打包脚本
├── config.py                       # 全局配置（窗口标题、版本号等）
├── requirements.txt                # Python依赖列表
├── chain_rules.json                # 连锁规则配置
│
├── svd_tool/                       # ★ 主代码目录
│   ├── __init__.py
│   ├── main.py                     # QApplication初始化
│   │
│   ├── core/                       # ★ 核心层（无UI依赖）
│   │   ├── data_model.py           #   数据模型（Field/Register/Peripheral/Interrupt/DeviceInfo）
│   │   ├── svd_parser.py           #   SVD文件解析器（DOM方式）
│   │   ├── chunked_svd_parser.py   #   SVD分块解析器（大文件优化）
│   │   ├── svd_generator.py        #   SVD文件生成器
│   │   ├── validators.py           #   输入验证器
│   │   ├── command_history.py      #   撤销/重做（Command模式）
│   │   └── constants.py            #   常量定义（访问权限选项等）
│   │
│   ├── ui/                         # ★ UI层
│   │   ├── main_window_refactored.py  # 主窗口（协调器角色）
│   │   ├── dialog_factories.py        # 对话框工厂（外设/寄存器/位域/中断编辑对话框）
│   │   ├── dialogs.py                 # 基础对话框（BaseEditDialog）
│   │   ├── tree_manager.py            # 外设树管理
│   │   │
│   │   ├── components/                # UI组件
│   │   │   ├── state_manager.py       #   状态管理器（核心枢纽，管理所有编辑状态）
│   │   │   ├── ui_updater.py          #   UI更新器（数据→控件同步）
│   │   │   ├── tab_builder.py         #   标签页构建器（创建基本信息/中断等标签页）
│   │   │   └── peripheral_manager.py  #   外设管理器
│   │   │
│   │   ├── managers/                  # 功能管理器
│   │   │   ├── file_operations.py     #   文件操作（打开/保存）
│   │   │   ├── interrupt_manager.py   #   中断管理器
│   │   │   └── device_info_manager.py #   设备信息管理器
│   │   │
│   │   └── widgets/                   # 自定义控件
│   │       ├── visualization_widget.py  # 可视化展示
│   │       ├── bit_field_widget.py      # 位域可视化
│   │       └── realtime_preview.py     # 实时XML预览
│   │
│   ├── i18n/                         # 国际化
│   │   ├── i18n.py                   #   翻译函数 t()
│   │   ├── zh_CN.py                  #   中文翻译
│   │   └── en_US.py                  #   英文翻译
│   │
│   └── utils/                        # 工具函数
│
├── tests/                           # 测试目录
└── docs/                            # 文档目录''')

    doc.add_heading('1.2 架构分层', level=2)
    doc.add_paragraph(
        '项目采用经典的三层架构：\n\n'
        '① Core层（svd_tool/core/）：纯数据逻辑，不依赖任何UI组件。\n'
        '   负责数据模型定义、SVD文件解析/生成、数据验证、命令历史。\n\n'
        '② UI层（svd_tool/ui/）：所有PyQt6相关的界面代码。\n'
        '   负责界面展示、用户交互、将用户操作转换为对Core层的调用。\n\n'
        '③ 入口层（run.py、main.py）：程序启动和初始化。'
    )

    doc.add_heading('1.3 核心调用链路', level=2)
    doc.add_paragraph('用户操作的主要调用链路：')
    add_code_block(doc, '''用户点击"保存"
  → main_window_refactored.py (事件处理)
    → file_operations.py (文件操作管理器)
      → state_manager.py (获取当前数据)
        → device_info (数据模型)
      → svd_generator.py (生成XML)
        → 写入文件''')

    add_code_block(doc, '''用户编辑中断
  → dialog_factories.py (弹出编辑对话框)
    → 用户填写表单
  → state_manager.py (执行命令)
    → command_history.py (记录命令，支持撤销)
    → 数据模型更新
    → _notify_state_change()
      → ui_updater.py (刷新UI)
      → tree_manager.py (刷新树)''')

    # --- 第2章 数据模型 ---
    doc.add_heading('第2章 数据模型详解', level=1)

    doc.add_heading('2.1 模型层次关系', level=2)
    add_code_block(doc, '''DeviceInfo（设备信息 —— 最顶层）
  │
  ├── vendor: str                    # 厂商名称
  ├── copyright: str                 # 版权信息
  ├── author: str                    # 作者
  ├── license: str                   # 许可证
  │
  ├── cpu: CPUInfo                   # CPU配置
  │     ├── name: str                #   如 CM0/CM3/CM4
  │     ├── revision: str            #   如 r0p0
  │     ├── endian: str              #   little / big
  │     ├── mpu_present: bool
  │     ├── fpu_present: bool
  │     ├── nvic_prio_bits: int
  │     └── vendor_systick_config: bool
  │
  ├── peripherals: Dict[str, Peripheral]  # 外设字典（按名称索引）
  │     │
  │     └── Peripheral（外设）
  │           ├── name: str
  │           ├── base_address: str
  │           ├── description: str
  │           ├── display_name: str
  │           ├── group_name: str
  │           ├── derived_from: str        # 继承的外设名
  │           ├── address_block: Dict      # {offset, size, usage}
  │           ├── registers: Dict[str, Register]  # 寄存器字典
  │           │     │
  │           │     └── Register（寄存器）
  │           │           ├── name, offset, description
  │           │           ├── size, access, reset_value
  │           │           └── fields: Dict[str, Field]  # 位域字典
  │           │                 │
  │           │                 └── Field（位域）
  │           │                       ├── name, description
  │           │                       ├── bit_offset, bit_width
  │           │                       ├── access, reset_value
  │           │                       └── enumerated_values: List[Dict]
  │           │
  │           └── interrupts: List[Dict]   # 外设级中断引用列表
  │                 # 每项: {name, value, description, peripheral}
  │
  └── interrupts: Dict[str, Interrupt]    # ★ 全局中断字典（按名称索引）
        │
        └── Interrupt（中断 —— 支持多外设共用）
              ├── name: str               # 中断名称
              ├── value: int              # 中断号
              ├── description: str
              ├── peripheral: str         # 主要关联外设（向后兼容）
              └── peripherals: List[str]  # ★ 所有关联外设列表''')

    doc.add_heading('2.2 关键设计说明', level=2)

    doc.add_heading('2.2.1 中断的双层存储', level=3)
    doc.add_paragraph(
        '中断数据存在两份：\n\n'
        '① 全局层：DeviceInfo.interrupts — Dict[str, Interrupt]\n'
        '   这是中断的权威数据源，每个中断有唯一的名称作为key。\n'
        '   Interrupt对象的peripherals列表记录了所有关联的外设。\n\n'
        '② 外设层：Peripheral.interrupts — List[Dict]\n'
        '   这是外设视角的中断引用列表，用于SVD文件生成。\n'
        '   当全局中断变更时，通过 _sync_interrupt_to_peripherals() 同步到外设层。\n\n'
        '同步机制：\n'
        '· 添加/修改/删除中断时，state_manager自动同步\n'
        '· _sync_all_peripheral_interrupts() 可以完全重建外设层的中断列表'
    )

    doc.add_heading('2.2.2 外设继承机制', level=3)
    doc.add_paragraph(
        'Peripheral的derived_from字段实现SVD标准的继承机制：\n'
        '· 继承外设自动获得基础外设的寄存器定义\n'
        '· 继承是单向的，修改基础外设不影响已继承的外设\n'
        '· 继承的外设可以覆盖基础外设的寄存器定义'
    )

    # --- 第3章 核心模块解析 ---
    doc.add_heading('第3章 核心模块解析', level=1)

    # 3.1 SVD解析器
    doc.add_heading('3.1 SVD解析器（svd_parser.py）', level=2)
    doc.add_paragraph('负责将SVD XML文件解析为内部数据模型。')

    doc.add_heading('3.1.1 解析流程', level=3)
    add_code_block(doc, '''parse_file(filepath)
  → 读取XML文件
  → minidom解析XML
  → _parse_device_info()        # 解析设备基本信息
  │    ├── 解析 <name>, <version>, <description>
  │    ├── 解析 <vendor>        # 厂商名称
  │    ├── _parse_comments()    # 解析注释中的 Copyright/Author/License
  │    └── _parse_cpu_info()    # 解析CPU配置
  → _parse_peripherals()        # 解析所有外设
  │    └── 对每个外设:
  │         ├── _parse_peripheral()    # 解析外设属性
  │         ├── _parse_registers()     # 解析寄存器
  │         ├── _parse_fields()        # 解析位域
  │         └── _parse_interrupts()    # 解析外设级中断
  → _collect_interrupts_to_device()    # ★ 合并同名中断到全局字典
  → 返回 DeviceInfo''')

    doc.add_heading('3.1.2 多外设共用中断的解析', level=3)
    doc.add_paragraph(
        '_collect_interrupts_to_device() 方法遍历所有外设的中断列表。'
        '当遇到已存在的同名中断时，将当前外设追加到已有中断的 peripherals 列表中，而非覆盖。'
    )
    add_code_block(doc, '''# 核心逻辑伪代码
for peripheral in peripherals:
    for irq in peripheral.interrupts:
        if irq.name in global_interrupts:
            # 同名中断已存在 → 追加外设关联
            if peripheral.name not in existing.peripherals:
                existing.peripherals.append(peripheral.name)
        else:
            # 新中断 → 创建并记录
            global_interrupts[irq.name] = Interrupt(
                name=irq.name, value=irq.value,
                peripherals=[peripheral.name]
            )''')

    doc.add_heading('3.1.3 XML注释解析', level=3)
    doc.add_paragraph(
        '_parse_comments() 方法解析XML文件中的注释节点，提取 Copyright、Author、License 信息。\n'
        '这些信息存储在XML注释中（非标准SVD元素），格式约定如下：'
    )
    add_code_block(doc, '''<!-- Copyright (c) 2026 SinOneMicroelectronics. -->
<!-- Author: SA Team -->
<!-- License: Apache-2.0 -->''')

    # 3.2 SVD生成器
    doc.add_heading('3.2 SVD生成器（svd_generator.py）', level=2)
    doc.add_paragraph('负责将内部数据模型生成SVD XML文件。')

    doc.add_heading('3.2.1 生成流程', level=3)
    add_code_block(doc, '''generate() → str
  → 创建 <device> 根元素
  → 添加基本信息（name, version, vendor, description...）
  → 添加 copyright/author/license 注释
  → _add_peripherals()
  │    └── 对每个外设:
  │         ├── _create_peripheral_element()
  │         │    ├── 添加 name, displayName, description, groupName
  │         │    ├── 添加 baseAddress, addressBlock
  │         │    ├── 添加中断（从 peripheral.interrupts 列表）
  │         │    └── 添加寄存器（从 peripheral.registers 字典）
  │         │         └── 对每个寄存器:
  │         │              ├── 添加 name, addressOffset, size, access...
  │         │              └── 添加位域 fields
  │         └── 添加到 <peripherals> 元素
  → 返回格式化的XML字符串''')

    # 3.3 状态管理器
    doc.add_heading('3.3 状态管理器（state_manager.py）', level=2)
    doc.add_paragraph(
        'StateManager是整个应用的枢纽，管理所有编辑状态和数据变更。'
        '它持有唯一的 DeviceInfo 实例，所有数据修改都通过它进行。'
    )

    doc.add_heading('3.3.1 核心职责', level=3)
    doc.add_paragraph('• 持有 DeviceInfo 实例（唯一数据源）', style='List Bullet')
    doc.add_paragraph('• 提供增删改查方法（外设/寄存器/位域/中断）', style='List Bullet')
    doc.add_paragraph('• 管理命令历史（撤销/重做）', style='List Bullet')
    doc.add_paragraph('• 通知UI更新（通过 _notify_state_change）', style='List Bullet')
    doc.add_paragraph('• 数据验证', style='List Bullet')

    doc.add_heading('3.3.2 通知机制', level=3)
    add_code_block(doc, '''# 状态变更通知链路
state_manager._notify_state_change()
  → 调用所有注册的回调函数
  → ui_updater.update_all()          # 刷新UI控件
  → tree_manager.update_tree()       # 刷新外设树
  → realtime_preview.update()        # 刷新XML预览''')

    doc.add_heading('3.3.3 中断管理方法', level=3)
    add_styled_table(doc,
        ['方法', '说明'],
        [
            ['add_interrupt(interrupt)', '添加中断，同步到关联外设'],
            ['update_interrupt(name, interrupt)', '更新中断（支持撤销），处理名称变更'],
            ['delete_interrupt(name)', '删除中断，从所有关联外设移除'],
            ['_sync_interrupt_to_peripherals(interrupt)', '将中断同步到所有关联外设的interrupts列表'],
            ['_remove_interrupt_from_peripheral(name, periph)', '从指定外设的中断列表中移除'],
            ['_sync_all_peripheral_interrupts()', '完全重建所有外设的中断列表'],
        ],
        col_widths=[6, 10]
    )

    # 3.4 命令历史
    doc.add_heading('3.4 命令历史（command_history.py）', level=2)
    doc.add_paragraph(
        '使用Command设计模式实现撤销/重做。每个可撤销的操作封装为Command对象：'
    )
    add_code_block(doc, '''class Command:
    execute: Callable    # 执行函数
    undo: Callable       # 撤销函数
    description: str     # 操作描述
    selection_before: dict  # 执行前的选中状态
    selection_after: dict   # 执行后的选中状态''')
    doc.add_paragraph('支持的撤销操作：外设/寄存器/位域/中断的添加、修改、删除、移动。')

    # 3.5 对话框工厂
    doc.add_heading('3.5 对话框工厂（dialog_factories.py）', level=2)
    doc.add_paragraph('DialogFactory 创建各类编辑对话框：')
    add_styled_table(doc,
        ['对话框类', '创建方法', '用途'],
        [
            ['PeripheralEditDialog', 'create_peripheral_dialog()', '外设编辑'],
            ['RegisterEditDialog', 'create_register_dialog()', '寄存器编辑'],
            ['FieldEditDialog', 'create_field_dialog()', '位域编辑'],
            ['InterruptEditDialog', 'create_interrupt_dialog()', '中断编辑（支持多外设选择）'],
        ],
        col_widths=[4, 5, 6]
    )
    doc.add_paragraph()
    doc.add_paragraph(
        '所有对话框继承 BaseEditDialog（dialogs.py），提供统一的布局和按钮管理。'
        '每个对话框实现 setup_form()、load_data()、validate_input()、collect_data() 四个核心方法。'
    )

    # 3.6 UI更新器
    doc.add_heading('3.6 UI更新器（ui_updater.py）', level=2)
    doc.add_paragraph(
        'UIUpdater 负责将数据模型的变更同步到UI控件。当 state_manager 发出通知时，'
        'UIUpdater 读取最新数据并更新对应的控件（QLineEdit.setText()、QCheckBox.setChecked() 等）。\n\n'
        '关键原则：数据模型是唯一的真相来源（Single Source of Truth），UI只是数据的投影。'
    )

    # 3.7 位域可视化控件
    doc.add_heading('3.7 位域可视化控件（bit_field_widget.py）', level=2)
    doc.add_paragraph(
        'BitFieldWidget 是自定义的QWidget，用于绘制寄存器位域图。它使用QPainter直接绘制。'
    )
    doc.add_paragraph('核心特性：', style='List Bullet')
    doc.add_paragraph('• 16色调色板（FIELD_COLORS）自动分配颜色给不同位域', style='List Bullet')
    doc.add_paragraph('• 保留位（RESERVED_COLOR = 灰色）标识未定义区域', style='List Bullet')
    doc.add_paragraph('• 智能标签显示：宽位域内嵌名称，窄位域使用外部斜线连接标签', style='List Bullet')
    doc.add_paragraph('• 鼠标跟踪：悬停高亮 + QToolTip详细信息', style='List Bullet')
    doc.add_paragraph('• 信号：field_clicked（选中位域）、jump_to_source_peripheral（跳转继承源）', style='List Bullet')
    doc.add_paragraph('• 调用 set_register_data(register, fields, source_peripheral) 更新数据', style='List Bullet')

    # 3.8 连锁规则引擎
    doc.add_heading('3.8 连锁规则引擎', level=2)
    doc.add_paragraph(
        '连锁操作允许在删除元素时自动触发关联元素的删除/修改。'
        '规则存储在 chain_rules.json 中，格式为 JSON。'
    )
    doc.add_paragraph('关键类：', style='List Bullet')
    doc.add_paragraph('• ChainRule：规则定义（源匹配条件 + 动作列表）', style='List Bullet')
    doc.add_paragraph('• ChainAction：单个动作（目标 + 动作类型）', style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('触发流程：')
    add_code_block(doc, '''delete_field() 被调用
  → 检查连锁功能是否全局启用
  → 遍历所有已启用规则
  → 对每条规则检查匹配：
     source_type == "field" ?
     source_peripheral 匹配？(支持通配符)
     source_register 匹配？
     source_field 匹配？
     trigger == "delete" ?
  → 对匹配的规则，执行每个 action：
     变量替换 ($PERIPHERAL → 实际名称)
     查找目标元素
     执行 delete/modify
  → 弹出对话框显示连锁操作结果''')
    doc.add_paragraph()
    doc.add_paragraph('扩展连锁规则：编辑 chain_rules.json 或通过菜单"工具→编辑连锁规则"在界面中编辑。')

    # --- 第4章 扩展指南 ---
    doc.add_heading('第4章 扩展与修改指南', level=1)

    doc.add_heading('4.1 添加新的数据字段', level=2)
    doc.add_paragraph('以添加"系列名称（series）"字段为例：')
    doc.add_paragraph()
    add_tip(doc, '步骤1：', '修改数据模型（data_model.py）')
    add_code_block(doc, '''# 在 DeviceInfo 类中添加字段
@dataclass
class DeviceInfo:
    name: str = ""
    version: str = "1.0"
    series: str = ""          # ← 新增字段
    ...''')

    add_tip(doc, '步骤2：', '修改解析器（svd_parser.py）')
    add_code_block(doc, '''# 在 _parse_device_info() 中添加解析
series_node = self._get_direct_child(device_node, "series")
if series_node and series_node.firstChild:
    self.device_info.series = series_node.firstChild.data.strip()''')

    add_tip(doc, '步骤3：', '修改生成器（svd_generator.py）')
    add_code_block(doc, '''# 在生成XML的方法中添加
if self.device_info.series:
    ET.SubElement(root, "series").text = self.device_info.series''')

    add_tip(doc, '步骤4：', '修改UI（tab_builder.py 添加控件，ui_updater.py 添加同步）')

    doc.add_heading('4.2 添加新的编辑对话框', level=2)
    doc.add_paragraph('1. 在 dialog_factories.py 中创建新类，继承 BaseEditDialog', style='List Bullet')
    doc.add_paragraph('2. 实现 setup_form()（创建控件）、load_data()（填充数据）、validate_input()（验证）、collect_data()（收集结果）', style='List Bullet')
    doc.add_paragraph('3. 在 DialogFactory 中添加创建方法', style='List Bullet')
    doc.add_paragraph('4. 在 state_manager.py 中添加对应的状态管理方法', style='List Bullet')

    doc.add_heading('4.3 添加新的SVD元素支持', level=2)
    doc.add_paragraph(
        'SVD标准还有很多元素尚未支持（如 enumeratedValues、addressBlock 多个等）。'
        '扩展步骤：'
    )
    doc.add_paragraph('1. 在 data_model.py 添加对应字段', style='List Bullet')
    doc.add_paragraph('2. 在 svd_parser.py 的 _parse_xxx() 方法中添加解析', style='List Bullet')
    doc.add_paragraph('3. 在 svd_generator.py 的 _create_xxx_element() 方法中添加生成', style='List Bullet')
    doc.add_paragraph('4. 在对话框中添加对应的编辑控件', style='List Bullet')
    doc.add_paragraph('5. 在 ui_updater.py 中添加UI同步逻辑', style='List Bullet')
    doc.add_paragraph('6. 编写测试用例验证', style='List Bullet')

    doc.add_heading('4.4 修改中断逻辑', level=2)
    doc.add_paragraph('中断相关代码分布在以下文件中：')
    add_styled_table(doc,
        ['文件', '职责', '关键方法/类'],
        [
            ['data_model.py', 'Interrupt 数据类定义', 'Interrupt（name/value/peripherals）'],
            ['svd_parser.py', '解析中断并合并同名', '_collect_interrupts_to_device()'],
            ['svd_generator.py', '写入中断到XML', '_add_interrupt_to_peripheral()'],
            ['state_manager.py', '中断状态管理和同步', 'add/update/delete_interrupt()'],
            ['interrupt_manager.py', '中断UI交互', 'InterruptManager'],
            ['dialog_factories.py', '中断编辑对话框', 'InterruptEditDialog'],
        ],
        col_widths=[4, 4, 7]
    )

    doc.add_heading('4.5 修改注意事项', level=2)
    doc.add_paragraph('• 修改数据模型后，必须同步修改解析器和生成器', style='List Bullet')
    doc.add_paragraph('• 添加新的撤销操作时，务必正确实现 undo 函数', style='List Bullet')
    doc.add_paragraph('• UI控件名称（objectName）必须与 layout_manager.get_widget() 的key一致', style='List Bullet')
    doc.add_paragraph('• 用户可见的文本使用 t() 函数包裹以支持国际化', style='List Bullet')
    doc.add_paragraph('• 修改后运行 pytest 验证测试通过', style='List Bullet')

    # --- 第5章 编码规范 ---
    doc.add_heading('第5章 编码规范', level=1)
    add_styled_table(doc,
        ['规范项', '要求'],
        [
            ['语言', 'Python 3.9+，使用类型注解（Type Hints）'],
            ['风格', '遵循 PEP 8'],
            ['命名', '类名 PascalCase，函数/变量 snake_case'],
            ['国际化', '所有用户可见文本使用 t("key") 函数'],
            ['UI访问', '通过 widget_manager.get_widget() 获取控件，不直接引用'],
            ['状态通知', '修改数据后调用 _notify_state_change()'],
            ['可撤销操作', '使用 Command 模式封装 execute/undo'],
            ['日志', '使用 logging 模块记录关键操作'],
            ['测试', '新增功能需在 tests/ 中添加测试用例'],
        ],
        col_widths=[4, 12]
    )

    # 保存
    path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'SVD_Editor_维护文档_v2.docx')
    doc.save(path)
    print(f'维护文档已生成: {path}')
    return path


if __name__ == '__main__':
    generate_user_manual()
    generate_maintenance_doc()
    print('\n两份文档均已生成完成！')